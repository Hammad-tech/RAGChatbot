####################################################################
#                         import
####################################################################

import streamlit as st

# Set page config must be the first Streamlit command
st.set_page_config(
    page_title="Chat With Your Data",
    page_icon="🤖",
    layout="wide"
)

import warnings
from dotenv import load_dotenv
import os
import glob
from pathlib import Path
from datetime import datetime
import sqlite3
from typing import Optional
import json

# Load environment variables
load_dotenv()

warnings.filterwarnings("ignore", category=FutureWarning)

# Initialize session state variables
if 'openai_api_key' not in st.session_state:
    st.session_state.openai_api_key = os.getenv("OPENAI_API_KEY", "")
if 'google_api_key' not in st.session_state:
    st.session_state.google_api_key = os.getenv("GOOGLE_API_KEY", "")
if 'hf_api_key' not in st.session_state:
    st.session_state.hf_api_key = os.getenv("HUGGINGFACEHUB_API_TOKEN", "")
if 'selected_model' not in st.session_state:
    st.session_state.selected_model = "gpt-3.5-turbo"
if 'temperature' not in st.session_state:
    st.session_state.temperature = 0.5
if 'top_p' not in st.session_state:
    st.session_state.top_p = 0.95
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'user_id' not in st.session_state:
    st.session_state.user_id = None
if 'credits' not in st.session_state:
    st.session_state.credits = 2.0  # Default 2 free credits
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None
if 'documents_loaded' not in st.session_state:
    st.session_state.documents_loaded = False
if 'LLM_provider' not in st.session_state:
    st.session_state.LLM_provider = "OpenAI"  # Default to OpenAI
if 'chain' not in st.session_state:
    st.session_state.chain = None
if 'memory' not in st.session_state:
    st.session_state.memory = None
if 'retriever' not in st.session_state:
    st.session_state.retriever = None
if 'pretrained_vector_store' not in st.session_state:
    st.session_state.pretrained_vector_store = None
if 'assistant_language' not in st.session_state:
    st.session_state.assistant_language = "english"
if 'retriever_type' not in st.session_state:
    st.session_state.retriever_type = "Vectorstore backed retriever"

# Import openai and google_genai as main LLM services
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_google_genai import GoogleGenerativeAIEmbeddings

# langchain prompts, memory, chains...
from langchain.prompts import PromptTemplate, ChatPromptTemplate
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory, ConversationSummaryBufferMemory

from langchain.schema import format_document

# document loaders
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    DirectoryLoader,
    CSVLoader,
    Docx2txtLoader,
)

# text_splitter
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    CharacterTextSplitter,
)

# OutputParser
from langchain_core.output_parsers import StrOutputParser

# Import chroma as the vector store
from langchain_community.vectorstores import Chroma

# Contextual_compression
from langchain.retrievers.document_compressors import DocumentCompressorPipeline
from langchain_community.document_transformers import (
    EmbeddingsRedundantFilter,
    LongContextReorder,
)
from langchain.retrievers.document_compressors import EmbeddingsFilter
from langchain.retrievers import ContextualCompressionRetriever

# HuggingFace
from langchain_community.embeddings import HuggingFaceInferenceAPIEmbeddings
from langchain_community.llms import HuggingFaceHub

####################################################################
#              Config: LLM services, assistant language,...
####################################################################
list_LLM_providers = [
    ":rainbow[**OpenAI**]",
    "**Google Generative AI**",
    ":hugging_face: **HuggingFace**",
]

dict_welcome_message = {
    "english": "How can I assist you today?",
    "french": "Comment puis-je vous aider aujourd'hui ?",
    "spanish": "¿Cómo puedo ayudarle hoy?",
    "german": "Wie kann ich Ihnen heute helfen?",
    "russian": "Чем я могу помочь вам сегодня?",
    "chinese": "我今天能帮你什么？",
    "arabic": "كيف يمكنني مساعدتك اليوم؟",
    "portuguese": "Como posso ajudá-lo hoje?",
    "italian": "Come posso assistervi oggi?",
    "Japanese": "今日はどのようなご用件でしょうか?",
}

list_retriever_types = [
    "Contextual compression",
    "Vectorstore backed retriever",
]

TMP_DIR = Path(__file__).resolve().parent.joinpath("data", "tmp")
LOCAL_VECTOR_STORE_DIR = (
    Path(__file__).resolve().parent.joinpath("data", "vector_stores")
)

# Add after the imports section
INITIAL_DATA_DIR = Path(__file__).resolve().parent.joinpath("data", "initial_training")
INITIAL_DATA_DIR.mkdir(parents=True, exist_ok=True)

# Add after the imports section
PRETRAINED_VECTOR_STORE_DIR = Path(__file__).resolve().parent.joinpath("data", "pretrained_vector_store")
PRETRAINED_VECTOR_STORE_DIR.mkdir(parents=True, exist_ok=True)

# Add after the imports section
DATABASE_DIR = Path(__file__).resolve().parent.joinpath("data", "database")
DATABASE_DIR.mkdir(parents=True, exist_ok=True)

# Admin credentials (in production, use environment variables)
ADMIN_USERNAME = "admin"
ADMIN_PASSWORD = "admin123"  # Change this in production
ADMIN_EMAIL = "admin@example.com"

####################################################################
#            Create app interface with streamlit
####################################################################
st.title("🤖 RAG chatbot")

# API keys
st.session_state.openai_api_key = ""
st.session_state.google_api_key = ""
st.session_state.hf_api_key = ""

# Create necessary directories
for dir_path in [
    Path(__file__).resolve().parent.joinpath("data"),
    Path(__file__).resolve().parent.joinpath("data", "database"),
    Path(__file__).resolve().parent.joinpath("data", "pretrained_vector_store"),
    Path(__file__).resolve().parent.joinpath("data", "tmp"),
    Path(__file__).resolve().parent.joinpath("data", "vector_stores")
]:
    dir_path.mkdir(parents=True, exist_ok=True)

# Update model list
list_models_openai = [
    "gpt-4-turbo-preview",
    "gpt-4",
    "gpt-3.5-turbo-0125",
    "gpt-3.5-turbo",
]

def expander_model_parameters(LLM_provider="OpenAI", text_input_API_key="OpenAI API Key - [Get an API key](https://platform.openai.com/account/api-keys)", list_models=list_models_openai):
    """Add a text_input (for API key) and a streamlit expander containing models and parameters."""
    st.session_state.LLM_provider = LLM_provider

    if LLM_provider == "OpenAI":
        env_api_key = os.getenv("OPENAI_API_KEY")
        if env_api_key:
            st.session_state.openai_api_key = env_api_key
        else:
            st.session_state.openai_api_key = st.text_input(text_input_API_key, type="password", placeholder="insert your API key")
        st.session_state.google_api_key = ""
        st.session_state.hf_api_key = ""

    if LLM_provider == "Google":
        env_api_key = os.getenv("GOOGLE_API_KEY")
        if env_api_key:
            st.session_state.google_api_key = env_api_key
        else:
            st.session_state.google_api_key = st.text_input(text_input_API_key, type="password", placeholder="insert your API key")
        st.session_state.openai_api_key = ""
        st.session_state.hf_api_key = ""

    if LLM_provider == "HuggingFace":
        env_api_key = os.getenv("HUGGINGFACEHUB_API_TOKEN")
        if env_api_key:
            st.session_state.hf_api_key = env_api_key
        else:
            st.session_state.hf_api_key = st.text_input(text_input_API_key, type="password", placeholder="insert your API key")
        st.session_state.openai_api_key = ""
        st.session_state.google_api_key = ""

    with st.expander("**Models and parameters**"):
        st.session_state.selected_model = st.selectbox(f"Choose {LLM_provider} model", list_models)
        st.session_state.temperature = st.slider("temperature", min_value=0.0, max_value=1.0, value=0.5, step=0.1)
        st.session_state.top_p = st.slider("top_p", min_value=0.0, max_value=1.0, value=0.95, step=0.05)


def load_initial_training_data():
    """Load initial training data from the initial_training directory"""
    try:
        # Load documents from initial training directory
        documents = []
        
        # Load PDF files
        pdf_loader = DirectoryLoader(
            INITIAL_DATA_DIR.as_posix(), 
            glob="**/*.pdf", 
            loader_cls=PyPDFLoader, 
            show_progress=True
        )
        documents.extend(pdf_loader.load())
        
        # Load TXT files
        txt_loader = DirectoryLoader(
            INITIAL_DATA_DIR.as_posix(), 
            glob="**/*.txt", 
            loader_cls=TextLoader, 
            show_progress=True
        )
        documents.extend(txt_loader.load())
        
        # Load CSV files
        csv_loader = DirectoryLoader(
            INITIAL_DATA_DIR.as_posix(), 
            glob="**/*.csv", 
            loader_cls=CSVLoader, 
            show_progress=True,
            loader_kwargs={"encoding":"utf8"}
        )
        documents.extend(csv_loader.load())
        
        # Load DOCX files
        doc_loader = DirectoryLoader(
            INITIAL_DATA_DIR.as_posix(),
            glob="**/*.docx",
            loader_cls=Docx2txtLoader,
            show_progress=True,
        )
        documents.extend(doc_loader.load())
        
        return documents
    except Exception as e:
        st.error(f"Error loading initial training data: {str(e)}")
        return []

def load_pretrained_vector_store():
    """Load the pre-trained vector store"""
    try:
        # Ensure we have an API key before proceeding
        if not st.session_state.openai_api_key and not st.session_state.google_api_key and not st.session_state.hf_api_key:
            st.error("Please provide an API key for your selected LLM provider")
            return False
            
        embeddings = select_embeddings_model()
        st.session_state.pretrained_vector_store = Chroma(
            embedding_function=embeddings,
            persist_directory=str(PRETRAINED_VECTOR_STORE_DIR)
        )
        
        # Create retriever for pre-trained data
        st.session_state.retriever = create_retriever(
            vector_store=st.session_state.pretrained_vector_store,
            embeddings=embeddings,
            retriever_type=st.session_state.retriever_type
        )
        
        # Create chain for pre-trained data
        st.session_state.chain, st.session_state.memory = create_ConversationalRetrievalChain(
            retriever=st.session_state.retriever,
            chain_type="stuff",
            language=st.session_state.assistant_language
        )
        
        return True
    except Exception as e:
        st.error(f"Error loading pre-trained vector store: {str(e)}")
        return False

def process_documents(documents):
    """Process documents and create vector store"""
    if not documents:
        st.error("No valid documents found")
        return False
    
    try:
        # Split documents into chunks
        chunks = split_documents_to_chunks(documents)
        
        # Create embeddings
        embeddings = select_embeddings_model()
        
        # Create vector store
        vector_store_name = f"vector_store_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        persist_directory = LOCAL_VECTOR_STORE_DIR / vector_store_name
        persist_directory.mkdir(parents=True, exist_ok=True)
        
        st.session_state.vector_store = Chroma.from_documents(
            documents=chunks,
            embedding=embeddings,
            persist_directory=str(persist_directory)
        )
        
        # Create retriever
        st.session_state.retriever = create_retriever(
            vector_store=st.session_state.vector_store,
            embeddings=embeddings,
            retriever_type=st.session_state.retriever_type
        )
        
        # Create chain
        st.session_state.chain, st.session_state.memory = create_ConversationalRetrievalChain(
            retriever=st.session_state.retriever,
            chain_type="stuff",
            language=st.session_state.assistant_language
        )
        
        st.session_state.documents_loaded = True
        st.success("Documents processed successfully! You can now chat with your data.")
        return True
        
    except Exception as e:
        st.error(f"Error processing documents: {str(e)}")
        return False

def process_uploaded_documents(uploaded_files):
    """Process uploaded documents and combine with pre-trained data"""
    if not uploaded_files:
        st.warning("Please upload at least one document")
        return False
    
    try:
        # Create temporary directory if it doesn't exist
        TMP_DIR.mkdir(parents=True, exist_ok=True)
        
        # Save uploaded files to temporary directory
        for uploaded_file in uploaded_files:
            file_path = TMP_DIR / uploaded_file.name
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
        
        # Load documents
        documents = langchain_document_loader()
        
        if not documents:
            st.error("No valid documents found in the uploaded files")
            return False
        
        # Split documents into chunks
        chunks = split_documents_to_chunks(documents)
        
        # Create embeddings
        embeddings = select_embeddings_model()
        
        # Create vector store for uploaded documents
        vector_store_name = f"user_vector_store_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        persist_directory = LOCAL_VECTOR_STORE_DIR / vector_store_name
        persist_directory.mkdir(parents=True, exist_ok=True)
        
        user_vector_store = Chroma.from_documents(
            documents=chunks,
            embedding=embeddings,
            persist_directory=str(persist_directory)
        )
        
        # Combine with pre-trained vector store
        if hasattr(st.session_state, 'pretrained_vector_store'):
            # Create a combined retriever
            base_retriever = user_vector_store.as_retriever()
            pretrained_retriever = st.session_state.pretrained_vector_store.as_retriever()
            
            # Create a combined retriever that searches both stores
            st.session_state.retriever = create_retriever(
                vector_store=user_vector_store,
                embeddings=embeddings,
                retriever_type=st.session_state.retriever_type,
                additional_retriever=pretrained_retriever
            )
        else:
            # If no pre-trained store, just use the user's vector store
            st.session_state.retriever = create_retriever(
                vector_store=user_vector_store,
                embeddings=embeddings,
                retriever_type=st.session_state.retriever_type
            )
        
        # Create chain
        st.session_state.chain, st.session_state.memory = create_ConversationalRetrievalChain(
            retriever=st.session_state.retriever,
            chain_type="stuff",
            language=st.session_state.assistant_language
        )
        
        st.session_state.documents_loaded = True
        st.success("Documents processed successfully! You can now chat with your data.")
        return True
        
    except Exception as e:
        st.error(f"Error processing documents: {str(e)}")
        return False
    finally:
        # Clean up temporary files
        delte_temp_files()

def sidebar_and_documentChooser():
    """Create the sidebar with document upload and model selection"""
    with st.sidebar:
        st.caption("🚀 A retrieval augmented generation chatbot powered by 🔗 Langchain, OpenAI, Google Generative AI and 🤗")
        st.write("")

        # Model selection first (needed for embeddings)
        st.subheader("Select Model")
        llm_chooser = st.radio("Select provider", list_LLM_providers, captions=[
            "[OpenAI pricing page](https://openai.com/pricing)",
            "Rate limit: 60 requests per minute.",
            "**Free access.**",
        ])

        if llm_chooser == list_LLM_providers[0]:
            expander_model_parameters(LLM_provider="OpenAI", text_input_API_key="OpenAI API Key - [Get an API key](https://platform.openai.com/account/api-keys)", list_models=list_models_openai)
        elif llm_chooser == list_LLM_providers[1]:
            expander_model_parameters(LLM_provider="Google", text_input_API_key="Google API Key - [Get an API key](https://makersuite.google.com/app/apikey)", list_models=["gemini-pro"])
        else:
            expander_model_parameters(LLM_provider="HuggingFace", text_input_API_key="HuggingFace API key - [Get an API key](https://huggingface.co/settings/tokens)", list_models=["mistralai/Mistral-7B-Instruct-v0.2"])

        # Assistant language
        st.write("")
        st.session_state.assistant_language = st.selectbox("Assistant language", list(dict_welcome_message.keys()))

        st.divider()
        st.subheader("Retrieval Settings")
        st.session_state.retriever_type = st.selectbox("Select retriever type", list_retriever_types)

        st.divider()
        
        # Document upload section
        st.subheader("Upload Documents")
        uploaded_files = st.file_uploader("Choose files", type=["pdf", "txt", "docx", "csv"], accept_multiple_files=True)
        
        if uploaded_files:
            st.write(f"Selected files: {[f.name for f in uploaded_files]}")
            if st.button("Process Documents"):
                with st.spinner("Processing documents..."):
                    if process_uploaded_documents(uploaded_files):
                        st.success("Documents processed successfully! You can now chat with your data.")
                    else:
                        st.error("Failed to process documents. Please check your files and try again.")
        elif not st.session_state.documents_loaded:
            if not hasattr(st.session_state, 'pretrained_vector_store') or st.session_state.pretrained_vector_store is None:
                with st.spinner("Loading pre-trained data..."):
                    if load_pretrained_vector_store():
                        st.success("Pre-trained data loaded successfully!")
                        st.session_state.documents_loaded = True
                    else:
                        st.error("Failed to load pre-trained data")


####################################################################
#        Process documents and create vectorstor (Chroma dB)
####################################################################
def delte_temp_files():
    """delete files from the './data/tmp' folder"""
    files = glob.glob(TMP_DIR.as_posix() + "/*")
    for f in files:
        try:
            os.remove(f)
        except:
            pass


def langchain_document_loader():
    """
    Crete documnet loaders for PDF, TXT and CSV files.
    https://python.langchain.com/docs/modules/data_connection/document_loaders/file_directory
    """

    documents = []

    txt_loader = DirectoryLoader(
        TMP_DIR.as_posix(), glob="**/*.txt", loader_cls=TextLoader, show_progress=True
    )
    documents.extend(txt_loader.load())

    pdf_loader = DirectoryLoader(
        TMP_DIR.as_posix(), glob="**/*.pdf", loader_cls=PyPDFLoader, show_progress=True
    )
    documents.extend(pdf_loader.load())

    csv_loader = DirectoryLoader(
        TMP_DIR.as_posix(), glob="**/*.csv", loader_cls=CSVLoader, show_progress=True,
        loader_kwargs={"encoding":"utf8"}
    )
    documents.extend(csv_loader.load())

    doc_loader = DirectoryLoader(
        TMP_DIR.as_posix(),
        glob="**/*.docx",
        loader_cls=Docx2txtLoader,
        show_progress=True,
    )
    documents.extend(doc_loader.load())
    return documents


def split_documents_to_chunks(documents):
    """Split documents to chunks using RecursiveCharacterTextSplitter."""

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1600, chunk_overlap=200)
    chunks = text_splitter.split_documents(documents)
    return chunks


def select_embeddings_model():
    """Select embeddings models: OpenAIEmbeddings or GoogleGenerativeAIEmbeddings."""
    if st.session_state.LLM_provider == "OpenAI":
        embeddings = OpenAIEmbeddings(api_key=st.session_state.openai_api_key)

    if st.session_state.LLM_provider == "Google":
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001", google_api_key=st.session_state.google_api_key
        )

    if st.session_state.LLM_provider == "HuggingFace":
        embeddings = HuggingFaceInferenceAPIEmbeddings(
            api_key=st.session_state.hf_api_key, model_name="thenlper/gte-large"
        )

    return embeddings


def create_retriever(
    vector_store,
    embeddings,
    retriever_type="Contextual compression",
    base_retriever_search_type="similarity",
    base_retriever_k=16,
    compression_retriever_k=20,
    additional_retriever=None
):
    """Create a retriever that can optionally combine with another retriever"""
    base_retriever = Vectorstore_backed_retriever(
        vectorstore=vector_store,
        search_type=base_retriever_search_type,
        k=base_retriever_k
    )
    
    if additional_retriever:
        # Create a combined retriever that searches both stores
        from langchain.retrievers import EnsembleRetriever
        ensemble_retriever = EnsembleRetriever(
            retrievers=[base_retriever, additional_retriever],
            weights=[0.5, 0.5]  # Equal weights for both retrievers
        )
        base_retriever = ensemble_retriever

    if retriever_type == "Vectorstore backed retriever":
        return base_retriever

    elif retriever_type == "Contextual compression":
        compression_retriever = create_compression_retriever(
            embeddings=embeddings,
            base_retriever=base_retriever,
            k=compression_retriever_k,
        )
        return compression_retriever


def Vectorstore_backed_retriever(
    vectorstore, search_type="similarity", k=4, score_threshold=None
):
    """create a vectorsore-backed retriever
    Parameters:
        search_type: Defines the type of search that the Retriever should perform.
            Can be "similarity" (default), "mmr", or "similarity_score_threshold"
        k: number of documents to return (Default: 4)
        score_threshold: Minimum relevance threshold for similarity_score_threshold (default=None)
    """
    search_kwargs = {}
    if k is not None:
        search_kwargs["k"] = k
    if score_threshold is not None:
        search_kwargs["score_threshold"] = score_threshold

    retriever = vectorstore.as_retriever(
        search_type=search_type, search_kwargs=search_kwargs
    )
    return retriever


def create_compression_retriever(
    embeddings, base_retriever, chunk_size=500, k=16, similarity_threshold=None
):
    """Build a ContextualCompressionRetriever.
    We wrap the the base_retriever (a Vectorstore-backed retriever) in a ContextualCompressionRetriever.
    The compressor here is a Document Compressor Pipeline, which splits documents
    to smaller chunks, removes redundant documents, filters the top relevant documents,
    and reorder the documents so that the most relevant are at beginning / end of the list.

    Parameters:
        embeddings: OpenAIEmbeddings or GoogleGenerativeAIEmbeddings.
        base_retriever: a Vectorstore-backed retriever.
        chunk_size (int): Docs will be splitted into smaller chunks using a CharacterTextSplitter with a default chunk_size of 500.
        k (int): top k relevant documents to the query are filtered using the EmbeddingsFilter. default =16.
        similarity_threshold : similarity_threshold of the  EmbeddingsFilter. default =None
    """

    # 1. splitting docs into smaller chunks
    splitter = CharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=0, separator=". "
    )

    # 2. removing redundant documents
    redundant_filter = EmbeddingsRedundantFilter(embeddings=embeddings)

    # 3. filtering based on relevance to the query
    relevant_filter = EmbeddingsFilter(
        embeddings=embeddings, k=k, similarity_threshold=similarity_threshold
    )

    # 4. Reorder the documents

    # Less relevant document will be at the middle of the list and more relevant elements at beginning / end.
    # Reference: https://python.langchain.com/docs/modules/data_connection/retrievers/long_context_reorder
    reordering = LongContextReorder()

    # 5. create compressor pipeline and retriever
    pipeline_compressor = DocumentCompressorPipeline(
        transformers=[splitter, redundant_filter, relevant_filter, reordering]
    )
    compression_retriever = ContextualCompressionRetriever(
        base_compressor=pipeline_compressor, base_retriever=base_retriever
    )

    return compression_retriever


def CohereRerank_retriever(
    base_retriever, cohere_api_key, cohere_model="rerank-multilingual-v2.0", top_n=10
):
    """Build a ContextualCompressionRetriever using CohereRerank endpoint to reorder the results
    based on relevance to the query.

    Parameters:
       base_retriever: a Vectorstore-backed retriever
       cohere_api_key: the Cohere API key
       cohere_model: the Cohere model, in ["rerank-multilingual-v2.0","rerank-english-v2.0"], default = "rerank-multilingual-v2.0"
       top_n: top n results returned by Cohere rerank. default = 10.
    """

    compressor = CohereRerank(
        cohere_api_key=cohere_api_key, model=cohere_model, top_n=top_n
    )

    retriever_Cohere = ContextualCompressionRetriever(
        base_compressor=compressor, base_retriever=base_retriever
    )
    return retriever_Cohere


def chain_RAG_blocks():
    """The RAG system is composed of:
    - 1. Retrieval: includes document loaders, text splitter, vectorstore and retriever.
    - 2. Memory.
    - 3. Converstaional Retreival chain.
    """
    with st.spinner("Creating vectorstore..."):
        # Check inputs
        error_messages = []
        if (
            not st.session_state.openai_api_key
            and not st.session_state.google_api_key
            and not st.session_state.hf_api_key
        ):
            error_messages.append(
                f"insert your {st.session_state.LLM_provider} API key"
            )

        if (
            st.session_state.retriever_type == list_retriever_types[0]
            and not st.session_state.cohere_api_key
        ):
            error_messages.append(f"insert your Cohere API key")
        if not st.session_state.uploaded_file_list:
            error_messages.append("select documents to upload")
        if st.session_state.vector_store_name == "":
            error_messages.append("provide a Vectorstore name")

        if len(error_messages) == 1:
            st.session_state.error_message = "Please " + error_messages[0] + "."
        elif len(error_messages) > 1:
            st.session_state.error_message = (
                "Please "
                + ", ".join(error_messages[:-1])
                + ", and "
                + error_messages[-1]
                + "."
            )
        else:
            st.session_state.error_message = ""
            try:
                # 1. Delete old temp files
                delte_temp_files()

                # 2. Upload selected documents to temp directory
                if st.session_state.uploaded_file_list is not None:
                    for uploaded_file in st.session_state.uploaded_file_list:
                        error_message = ""
                        try:
                            temp_file_path = os.path.join(
                                TMP_DIR.as_posix(), uploaded_file.name
                            )
                            with open(temp_file_path, "wb") as temp_file:
                                temp_file.write(uploaded_file.read())
                        except Exception as e:
                            error_message += e
                    if error_message != "":
                        st.warning(f"Errors: {error_message}")

                    # 3. Load documents with Langchain loaders
                    documents = langchain_document_loader()

                    # 4. Split documents to chunks
                    chunks = split_documents_to_chunks(documents)
                    # 5. Embeddings
                    embeddings = select_embeddings_model()

                    # 6. Create a vectorstore
                    persist_directory = (
                        LOCAL_VECTOR_STORE_DIR.as_posix()
                        + "/"
                        + st.session_state.vector_store_name
                    )

                    try:
                        st.session_state.vector_store = Chroma.from_documents(
                            documents=chunks,
                            embedding=embeddings,
                            persist_directory=persist_directory,
                        )
                        st.info(
                            f"Vectorstore **{st.session_state.vector_store_name}** is created succussfully."
                        )

                        # 7. Create retriever
                        st.session_state.retriever = create_retriever(
                            vector_store=st.session_state.vector_store,
                            embeddings=embeddings,
                            retriever_type=st.session_state.retriever_type,
                            base_retriever_search_type="similarity",
                            base_retriever_k=16,
                            compression_retriever_k=20,
                            cohere_api_key=st.session_state.cohere_api_key,
                            cohere_model="rerank-multilingual-v2.0",
                            cohere_top_n=10,
                        )

                        # 8. Create memory and ConversationalRetrievalChain
                        (
                            st.session_state.chain,
                            st.session_state.memory,
                        ) = create_ConversationalRetrievalChain(
                            retriever=st.session_state.retriever,
                            chain_type="stuff",
                            language=st.session_state.assistant_language,
                        )

                        # 9. Cclear chat_history
                        clear_chat_history()

                    except Exception as e:
                        st.error(e)

            except Exception as error:
                st.error(f"An error occurred: {error}")


####################################################################
#                       Create memory
####################################################################


def create_memory(model_name="gpt-3.5-turbo", memory_max_token=None):
    """Creates a ConversationSummaryBufferMemory for gpt-3.5-turbo
    Creates a ConversationBufferMemory for the other models"""

    if model_name == "gpt-3.5-turbo":
        if memory_max_token is None:
            memory_max_token = 1024  # max_tokens for 'gpt-3.5-turbo' = 4096
        memory = ConversationSummaryBufferMemory(
            max_token_limit=memory_max_token,
            llm=ChatOpenAI(
                model_name="gpt-3.5-turbo",
                openai_api_key=st.session_state.openai_api_key,
                temperature=0.1,
            ),
            return_messages=True,
            memory_key="chat_history",
            output_key="answer",
            input_key="question",
        )
    else:
        memory = ConversationBufferMemory(
            return_messages=True,
            memory_key="chat_history",
            output_key="answer",
            input_key="question",
        )
    return memory


####################################################################
#          Create ConversationalRetrievalChain with memory
####################################################################


def answer_template(language="english"):
    """Pass the standalone question along with the chat history and context
    to the `LLM` wihch will answer."""

    template = f"""You are a helpful AI assistant. Your task is to answer questions based ONLY on the provided context (documents).
If the answer cannot be found in the context, say "I don't have enough information in my documents to answer this question."
Do not make up information or use external knowledge.

<context>
{{chat_history}}

{{context}} 
</context>

Question: {{question}}

Language: {language}.

Instructions:
1. Answer based ONLY on the provided context
2. If the answer is not in the context, say so
3. Be precise and concise
4. If relevant, cite the source document
5. Answer in the specified language
"""
    return template


def create_ConversationalRetrievalChain(
    retriever,
    chain_type="stuff",
    language="english",
):
    """Create a ConversationalRetrievalChain.
    First, it passes the follow-up question along with the chat history to an LLM which rephrases
    the question and generates a standalone query.
    This query is then sent to the retriever, which fetches relevant documents (context)
    and passes them along with the standalone question and chat history to an LLM to answer.
    """

    # 1. Define the standalone_question prompt.
    # Pass the follow-up question along with the chat history to the `condense_question_llm`
    # which rephrases the question and generates a standalone question.

    condense_question_prompt = PromptTemplate(
        input_variables=["chat_history", "question"],
        template="""Given the following conversation and a follow up question, 
rephrase the follow up question to be a standalone question, in its original language.\n\n
Chat History:\n{chat_history}\n
Follow Up Input: {question}\n
Standalone question:""",
    )

    # 2. Define the answer_prompt
    # Pass the standalone question + the chat history + the context (retrieved documents)
    # to the `LLM` wihch will answer

    answer_prompt = ChatPromptTemplate.from_template(answer_template(language=language))

    # 3. Add ConversationSummaryBufferMemory for gpt-3.5, and ConversationBufferMemory for the other models
    memory = create_memory(st.session_state.selected_model)

    # 4. Instantiate LLMs: standalone_query_generation_llm & response_generation_llm
    if st.session_state.LLM_provider == "OpenAI":
        standalone_query_generation_llm = ChatOpenAI(
            api_key=st.session_state.openai_api_key,
            model=st.session_state.selected_model,
            temperature=0.1,
        )
        response_generation_llm = ChatOpenAI(
            api_key=st.session_state.openai_api_key,
            model=st.session_state.selected_model,
            temperature=st.session_state.temperature,
            model_kwargs={"top_p": st.session_state.top_p},
        )
    if st.session_state.LLM_provider == "Google":
        standalone_query_generation_llm = ChatGoogleGenerativeAI(
            google_api_key=st.session_state.google_api_key,
            model=st.session_state.selected_model,
            temperature=0.1,
            convert_system_message_to_human=True,
        )
        response_generation_llm = ChatGoogleGenerativeAI(
            google_api_key=st.session_state.google_api_key,
            model=st.session_state.selected_model,
            temperature=st.session_state.temperature,
            top_p=st.session_state.top_p,
            convert_system_message_to_human=True,
        )

    if st.session_state.LLM_provider == "HuggingFace":
        standalone_query_generation_llm = HuggingFaceHub(
            repo_id=st.session_state.selected_model,
            huggingfacehub_api_token=st.session_state.hf_api_key,
            model_kwargs={
                "temperature": 0.1,
                "top_p": 0.95,
                "do_sample": True,
                "max_new_tokens": 1024,
            },
        )
        response_generation_llm = HuggingFaceHub(
            repo_id=st.session_state.selected_model,
            huggingfacehub_api_token=st.session_state.hf_api_key,
            model_kwargs={
                "temperature": st.session_state.temperature,
                "top_p": st.session_state.top_p,
                "do_sample": True,
                "max_new_tokens": 1024,
            },
        )

    # 5. Create the ConversationalRetrievalChain

    chain = ConversationalRetrievalChain.from_llm(
        condense_question_prompt=condense_question_prompt,
        combine_docs_chain_kwargs={"prompt": answer_prompt},
        condense_question_llm=standalone_query_generation_llm,
        llm=response_generation_llm,
        memory=memory,
        retriever=retriever,
        chain_type=chain_type,
        verbose=False,
        return_source_documents=True,
    )

    return chain, memory


def clear_chat_history():
    """clear chat history and memory."""
    # 1. re-initialize messages
    st.session_state.messages = [
        {
            "role": "assistant",
            "content": dict_welcome_message[st.session_state.assistant_language],
        }
    ]
    # 2. Clear memory (history)
    try:
        st.session_state.memory.clear()
    except:
        pass


def get_response_from_LLM(prompt):
    """Get response from LLM using the loaded documents"""
    if not st.session_state.documents_loaded:
        st.error("Please upload and process documents first")
        return None
        
    try:
        with st.spinner("Generating response..."):
            response = st.session_state.chain.invoke({"question": prompt})
            answer = response["answer"]
            save_chat_history(st.session_state.user_id, prompt, answer)
            st.session_state.messages.append({"role": "user", "content": prompt})
            st.session_state.messages.append({"role": "assistant", "content": answer})
            
            if "source_documents" in response and response["source_documents"]:
                with st.expander("Source Documents"):
                    for doc in response["source_documents"]:
                        st.write(f"Source: {doc.metadata.get('source', 'Unknown')}")
                        st.write(doc.page_content)
                        st.divider()
            
            return answer
    except Exception as e:
        st.error(f"Error getting response: {str(e)}")
        return None


####################################################################
#                         Chatbot
####################################################################
def chatbot():
    """Main chatbot function with authentication and credit management"""
    st.title("🤖 RAG chatbot")

    # Authentication
    if not st.session_state.user_id:
        st.sidebar.title("Authentication")
        auth_type = st.sidebar.radio("Choose action", ["Login", "Register"])
        
        if auth_type == "Login":
            username = st.sidebar.text_input("Username")
            password = st.sidebar.text_input("Password", type="password")
            if st.sidebar.button("Login"):
                user_id = login_user(username, password)
                if user_id:
                    st.sidebar.success("Logged in successfully!")
                    st.rerun()
                else:
                    st.sidebar.error("Invalid credentials")
        
        else:  # Register
            username = st.sidebar.text_input("Username")
            email = st.sidebar.text_input("Email")
            password = st.sidebar.text_input("Password", type="password")
            if st.sidebar.button("Register"):
                if register_user(username, email, password):
                    st.sidebar.success("Registration successful! Please login.")
                else:
                    st.sidebar.error("Username or email already exists")
        
        st.stop()
    
    # Show user info and credits
    st.sidebar.title("User Info")
    if st.session_state.is_admin:
        st.sidebar.write("👑 Admin Account")
        st.sidebar.write("Credits: Unlimited")
    else:
        st.sidebar.write(f"Credits remaining: {st.session_state.credits}")
    
    if st.sidebar.button("Logout"):
        st.session_state.user_id = None
        st.session_state.credits = 2.0
        st.session_state.is_admin = False
        st.rerun()

    # Admin controls
    if st.session_state.is_admin:
        st.sidebar.divider()
        st.sidebar.subheader("Admin Controls")
        if st.sidebar.button("Process Database Documents"):
            with st.spinner("Processing database documents..."):
                if process_database_documents():
                    st.sidebar.success("Database documents processed successfully!")
                else:
                    st.sidebar.error("Failed to process database documents")

    # Main chat interface
    sidebar_and_documentChooser()
    st.divider()
    
    col1, col2 = st.columns([7, 3])
    with col1:
        st.subheader("Chat with your data")
    with col2:
        st.button("Clear Chat History", on_click=clear_chat_history)

    # Initialize messages if not present
    if "messages" not in st.session_state:
        st.session_state.messages = [
            {
                "role": "assistant",
                "content": dict_welcome_message[st.session_state.assistant_language],
            }
        ]

    # Display chat messages
    for msg in st.session_state.messages:
        st.chat_message(msg["role"]).write(msg["content"])

    # Chat input
    if prompt := st.chat_input():
        if not st.session_state.documents_loaded:
            st.error("Please upload and process documents first")
            st.stop()
            
        if not st.session_state.is_admin and st.session_state.credits <= 0:
            st.error("Insufficient credits. Please purchase more credits to continue.")
            st.stop()
        
        if not st.session_state.openai_api_key and not st.session_state.google_api_key and not st.session_state.hf_api_key:
            st.info(f"Please insert your {st.session_state.LLM_provider} API key to continue.")
            st.stop()
        
        # Add user message to chat
        st.session_state.messages.append({"role": "user", "content": prompt})
        st.chat_message("user").write(prompt)
        
        # Get response from LLM
        with st.spinner("Thinking..."):
            response = get_response_from_LLM(prompt)
            
            if response:
                # Add assistant response to chat
                st.session_state.messages.append({"role": "assistant", "content": response})
                st.chat_message("assistant").write(response)
                
                # Only deduct credits for non-admin users
                if not st.session_state.is_admin:
                    st.session_state.credits -= 1
                    update_credits(st.session_state.user_id, st.session_state.credits)


####################################################################
#                         Database Functions
####################################################################

def init_db():
    """Initialize the SQLite database"""
    conn = sqlite3.connect('data/users.db')
    c = conn.cursor()
    
    # Drop existing tables to recreate with correct schema
    c.execute("DROP TABLE IF EXISTS users")
    c.execute("DROP TABLE IF EXISTS chat_history")
    
    # Create users table with all required columns
    c.execute('''CREATE TABLE IF NOT EXISTS users
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  username TEXT UNIQUE NOT NULL,
                  email TEXT UNIQUE NOT NULL,
                  password_hash TEXT NOT NULL,
                  credits REAL DEFAULT 2.0,
                  is_admin BOOLEAN DEFAULT 0,
                  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
    
    # Create chat_history table
    c.execute('''CREATE TABLE IF NOT EXISTS chat_history
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  user_id INTEGER NOT NULL,
                  question TEXT NOT NULL,
                  answer TEXT NOT NULL,
                  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                  FOREIGN KEY (user_id) REFERENCES users (id))''')
    
    # Create admin user if not exists
    c.execute("SELECT * FROM users WHERE username = ?", (ADMIN_USERNAME,))
    if not c.fetchone():
        c.execute("INSERT INTO users (username, email, password_hash, credits, is_admin) VALUES (?, ?, ?, ?, ?)",
                 (ADMIN_USERNAME, ADMIN_EMAIL, ADMIN_PASSWORD, float('inf'), 1))
    
    conn.commit()
    conn.close()

def register_user(username: str, email: str, password: str) -> bool:
    """Register a new user"""
    try:
        conn = sqlite3.connect('data/users.db')
        c = conn.cursor()
        
        # Check if username or email already exists
        c.execute("SELECT * FROM users WHERE username = ? OR email = ?", (username, email))
        if c.fetchone():
            return False
        
        # Hash password (in production, use proper password hashing)
        password_hash = password  # In production, use proper hashing
        
        # Insert new user
        c.execute("INSERT INTO users (username, email, password_hash) VALUES (?, ?, ?)",
                 (username, email, password_hash))
        
        conn.commit()
        conn.close()
        return True
    except:
        return False

def login_user(username: str, password: str) -> Optional[int]:
    """Login a user and return their user_id"""
    try:
        conn = sqlite3.connect('data/users.db')
        c = conn.cursor()
        
        # For admin, use direct password comparison
        if username == ADMIN_USERNAME:
            if password == ADMIN_PASSWORD:
                st.session_state.user_id = 1  # Admin user_id is 1
                st.session_state.credits = float('inf')
                st.session_state.is_admin = True
                return 1
            return None
            
        # For regular users, check username and password
        c.execute("SELECT id, password_hash, credits, is_admin FROM users WHERE username = ?", (username,))
        result = c.fetchone()
        
        if result and result[1] == password:  # Direct password comparison
            st.session_state.user_id = result[0]
            st.session_state.credits = result[2]
            st.session_state.is_admin = bool(result[3])
            return result[0]
        return None
    except:
        return None

def update_credits(user_id: int, credits: float) -> bool:
    """Update user credits"""
    try:
        conn = sqlite3.connect('data/users.db')
        c = conn.cursor()
        
        c.execute("UPDATE users SET credits = ? WHERE id = ?", (credits, user_id))
        conn.commit()
        conn.close()
        return True
    except:
        return False

def save_chat_history(user_id: int, question: str, answer: str) -> bool:
    """Save chat history"""
    try:
        conn = sqlite3.connect('data/users.db')
        c = conn.cursor()
        
        c.execute("INSERT INTO chat_history (user_id, question, answer) VALUES (?, ?, ?)",
                 (user_id, question, answer))
        
        conn.commit()
        conn.close()
        return True
    except:
        return False

def download_chat_history():
    """Download chat history as a document"""
    try:
        conn = sqlite3.connect('data/users.db')
        c = conn.cursor()
        
        # Get chat history for current user
        c.execute("""
            SELECT question, answer, created_at 
            FROM chat_history 
            WHERE user_id = ? 
            ORDER BY created_at DESC
        """, (st.session_state.user_id,))
        
        chat_history = c.fetchall()
        conn.close()
        
        if not chat_history:
            st.warning("No chat history found")
            return
        
        # Create document content
        doc_content = f"Chat History for User ID: {st.session_state.user_id}\n\n"
        for question, answer, timestamp in chat_history:
            doc_content += f"Timestamp: {timestamp}\n"
            doc_content += f"Question: {question}\n"
            doc_content += f"Answer: {answer}\n"
            doc_content += "-" * 50 + "\n\n"
        
        # Create download buttons for both TXT and DOCX formats
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="Download as TXT",
                data=doc_content,
                file_name=f"chat_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain"
            )
        with col2:
            # Create a Word document
            from docx import Document
            doc = Document()
            doc.add_heading(f'Chat History for User ID: {st.session_state.user_id}', 0)
            
            for question, answer, timestamp in chat_history:
                doc.add_paragraph(f'Timestamp: {timestamp}')
                doc.add_paragraph(f'Question: {question}')
                doc.add_paragraph(f'Answer: {answer}')
                doc.add_paragraph('-' * 50)
            
            # Save the document to a temporary file
            temp_file = "temp_chat_history.docx"
            doc.save(temp_file)
            
            # Read the file and create download button
            with open(temp_file, "rb") as f:
                st.download_button(
                    label="Download as DOCX",
                    data=f,
                    file_name=f"chat_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            # Clean up the temporary file
            import os
            os.remove(temp_file)
    except Exception as e:
        st.error(f"Error downloading chat history: {str(e)}")

def process_database_documents():
    """Process all documents in the database directory and create embeddings"""
    try:
        # Load all documents from database directory
        documents = []
        
        # Load PDF files
        pdf_loader = DirectoryLoader(
            DATABASE_DIR.as_posix(), 
            glob="**/*.pdf", 
            loader_cls=PyPDFLoader, 
            show_progress=True
        )
        documents.extend(pdf_loader.load())
        
        if not documents:
            st.error("No documents found in database directory")
            return False
        
        # Split documents into chunks
        chunks = split_documents_to_chunks(documents)
        
        # Create embeddings
        embeddings = select_embeddings_model()
        
        # Create vector store
        st.session_state.pretrained_vector_store = Chroma.from_documents(
            documents=chunks,
            embedding=embeddings,
            persist_directory=str(PRETRAINED_VECTOR_STORE_DIR)
        )
        
        # Create retriever
        st.session_state.retriever = create_retriever(
            vector_store=st.session_state.pretrained_vector_store,
            embeddings=embeddings,
            retriever_type=st.session_state.retriever_type
        )
        
        # Create chain
        st.session_state.chain, st.session_state.memory = create_ConversationalRetrievalChain(
            retriever=st.session_state.retriever,
            chain_type="stuff",
            language=st.session_state.assistant_language
        )
        
        st.session_state.documents_loaded = True
        return True
        
    except Exception as e:
        st.error(f"Error processing database documents: {str(e)}")
        return False

# Initialize database
init_db()

# Add download button to the sidebar
if st.session_state.user_id:
    st.sidebar.divider()
    download_chat_history()

if __name__ == "__main__":
    chatbot()
